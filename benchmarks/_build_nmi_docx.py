"""Build professionally formatted NMI paper as DOCX."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style definitions ──
style = doc.styles['Normal']
style.font.name = 'Garamond'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level, (size, bold, color, space_before) in enumerate([
    (18, True, RGBColor(0x1a, 0x23, 0x32), 24),   # Heading 1
    (14, True, RGBColor(0x2c, 0x3e, 0x50), 18),   # Heading 2
    (12, True, RGBColor(0x34, 0x49, 0x5e), 14),   # Heading 3
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Garamond'
    h.font.size = Pt(size)
    h.font.bold = bold
    h.font.color.rgb = color
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(8)
    h.paragraph_format.keep_with_next = True


def add_para(text, style='Normal', bold=False, italic=False, font_size=None, alignment=None, space_after=None, space_before=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if font_size:
        run.font.size = Pt(font_size)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def add_rich_para(segments, style='Normal', alignment=None, space_after=None):
    """Add paragraph with mixed formatting. segments = list of (text, {bold, italic, font_name, font_size, color})."""
    p = doc.add_paragraph(style=style)
    for text, fmt in segments:
        run = p.add_run(text)
        if fmt.get('bold'):
            run.bold = True
        if fmt.get('italic'):
            run.italic = True
        if fmt.get('font_name'):
            run.font.name = fmt['font_name']
        if fmt.get('font_size'):
            run.font.size = Pt(fmt['font_size'])
        if fmt.get('color'):
            run.font.color.rgb = fmt['color']
        if fmt.get('superscript'):
            run.font.superscript = True
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Garamond'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1a2332')
        cell.paragraphs[0]._element.get_or_add_pPr().append(shading)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Garamond'
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternate row shading
            if ri % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'f4f6f8')
                p._element.get_or_add_pPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_code(text):
    """Add code-formatted text."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════

doc.add_paragraph()  # spacing
doc.add_paragraph()

add_para(
    'Systematic Gauge Invariance Failure\nin Large Language Model Moral Cognition',
    bold=True, font_size=22, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
)

add_para(
    'Andrew H. Bond¹ ² and Lucas Thiele³',
    font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
)

add_rich_para([
    ('¹ Department of Computer Engineering, San Jose State University, San Jose, CA\n', {'font_size': 9, 'color': RGBColor(0x66, 0x66, 0x66)}),
    ('² Senior Member, IEEE\n', {'font_size': 9, 'color': RGBColor(0x66, 0x66, 0x66)}),
    ('³ Department of Cognitive Science, University of California, Los Angeles, CA', {'font_size': 9, 'color': RGBColor(0x66, 0x66, 0x66)}),
], alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Abstract', level=2)

add_para(
    'Large language models are increasingly deployed in contexts requiring moral judgment — content moderation, '
    'clinical decision support, legal analysis — yet our understanding of how their moral cognition responds to '
    'surface-level variation remains rudimentary. We introduce a geometric evaluation framework that treats moral '
    'judgment as a point in a multi-dimensional harm space and measures displacement under perturbations that should '
    'leave the evaluation unchanged. Applying this framework across five cognitive domains — learning, metacognition, '
    'attention, executive functions, and social cognition — we find that LLMs systematically fail gauge invariance: '
    'surface-level perturbations orthogonal to moral content (social pressure, emotional tone, linguistic framing, '
    'irrelevant sensory detail) displace moral judgments at 4.6–13.3 sigma significance across 5 model architectures. '
    'Critically, these five failure modes converge on a single geometric property: the model\'s reasoning process does '
    'not maintain separation between moral content and surface presentation. We identify a shared metacognitive '
    'intervention ceiling (~38% recovery across perturbation types), a sycophancy gradient that scales inversely with '
    'model stability, and architectural dissociations proving these are independent cognitive faculties that no single '
    'robustness score can capture. These findings demonstrate that moral cognition in current LLMs possesses some '
    'geometric symmetries (gender invariance, evaluation-order independence) but lacks others (framing invariance, '
    'emotional invariance, social-pressure invariance), providing a structured diagnostic for alignment that scalar '
    'benchmarks cannot deliver.',
    italic=True, font_size=10, space_after=18
)

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('1. Introduction', level=1)

add_para(
    'When a language model judges a moral scenario, its output can be understood as a point in a '
    'multi-dimensional space defined by moral dimensions such as physical harm, emotional harm, autonomy violation, '
    'and trust breach. A morally competent system should produce the same evaluation regardless of how the scenario '
    'is described — whether the language is euphemistic or dramatic, whether irrelevant sensory details are present, '
    'whether someone pressures the system to change its mind. This property — that morally equivalent inputs under '
    'different surface presentations should yield identical outputs — is what we term gauge invariance, borrowing '
    'from the physics of field theories where observable quantities must be invariant under admissible '
    're-descriptions of the coordinate system¹.'
)

add_para(
    'The question of whether LLMs possess this property has been approached piecewise: studies of sycophancy²˒³, '
    'framing effects⁴, anchoring bias⁵, and prompt sensitivity⁶ each probe one dimension of the problem. But these '
    'studies share a structural limitation: each measures a single perturbation type in isolation, producing scalar '
    'robustness scores that cannot distinguish models with qualitatively different vulnerability profiles. The Scalar '
    'Irrecoverability Theorem¹ formalizes this limitation: for a system evaluated across n independent dimensions of '
    'variation, any scalar summary destroys n − 1 directions of information. No post hoc procedure can recover the '
    'lost structure.'
)

add_para(
    'This paper takes a different approach. We define a shared geometric coordinate system — a 7-dimensional harm '
    'space — and apply five qualitatively distinct perturbation types across five cognitive domains, measuring '
    'displacement of the judgment vector in each case. The perturbation types (social pressure, emotional reframing, '
    'linguistic register, irrelevant sensory detail, and self-knowledge accuracy) correspond to gauge transformations '
    'that should leave the moral evaluation invariant. The cognitive domains (learning, metacognition, attention, '
    'executive functions, social cognition) correspond to distinct faculties that process these perturbations.'
)

add_para(
    'Our central finding is that gauge invariance fails systematically but selectively. Models are invariant under '
    'some transformations (gender swap, evaluation order) but not others (framing, emotional tone, social pressure). '
    'The pattern of which symmetries are preserved and which are broken constitutes a cognitive profile of each model '
    'architecture — a structured diagnostic that no single benchmark score can capture. Moreover, the five failure '
    'modes converge quantitatively: recovery rates plateau at approximately 38% across perturbation types, suggesting '
    'a shared metacognitive intervention ceiling, and the sycophancy gradient across model generations reveals that '
    'newer, more capable models are progressively worse at distinguishing valid from invalid corrections.'
)

# 1.1 Related Work
doc.add_heading('1.1 Related Work', level=2)

add_rich_para([
    ('Sycophancy and social pressure. ', {'bold': True}),
    ('Sharma et al.² demonstrated that LLMs shift responses to align with user opinions. Perez et al.³ developed '
     'model-written evaluations revealing consistent sycophantic tendencies. Our L2 protocol extends this work with a '
     'multi-turn design that establishes a committed baseline before applying corrections, operationalizing the '
     'intensity-zero identity principle.', {}),
])

add_rich_para([
    ('Framing and anchoring. ', {'bold': True}),
    ('Tversky and Kahneman⁴ established framing effects in human decision-making. Recent work has identified '
     'analogous effects in LLMs⁵˒⁶, but without the controlled separation of content from presentation that our '
     'three-tier data architecture provides.', {}),
])

add_rich_para([
    ('Calibration. ', {'bold': True}),
    ('Kadavath et al.⁷ showed that LLMs "mostly know what they know." Our M1 results challenge this conclusion in '
     'the moral domain: Fisher-combined miscalibration across 4 models reaches 9.3 sigma, with every model tested '
     'showing systematic overconfidence.', {}),
])

add_rich_para([
    ('Geometric and multi-dimensional evaluation. ', {'bold': True}),
    ('Bond¹ introduced the geometric ethics framework, proposing that moral situations occupy a differentiable '
     'manifold with tensor structure. Thiele⁸ demonstrated that the nine moral dimensions of this framework are '
     'linearly decodable from LaBSE representations, with moral and language signals occupying largely orthogonal '
     'subspaces. Our work extends this empirical program from representation probing to behavioral evaluation across '
     'five cognitive domains.', {}),
])

add_rich_para([
    ('Executive functions in AI. ', {'bold': True}),
    ('Diamond⁹ defined executive functions as cognitive flexibility, inhibitory control, and working memory. We '
     'operationalize these constructs for LLM evaluation, measuring framework switching, emotional anchoring '
     'resistance, counterfactual sensitivity, and party-tracking under complexity scaling.', {}),
])


# ═══════════════════════════════════════════════════════════════
# 2. FRAMEWORK AND METHODS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('2. Framework and Methods', level=1)

doc.add_heading('2.1 The Judgment Manifold', level=2)

add_para(
    'We model moral judgment as a mapping from scenario descriptions to points in a 7-dimensional harm space '
    'H = (physical, emotional, financial, autonomy, trust, social_impact, identity), where each dimension is scored '
    '0–10, yielding a total harm score in [0, 70]. This coordinate system is shared across all five benchmark tracks, '
    'ensuring that displacement measurements are commensurable.'
)

add_para(
    'A gauge transformation is any re-description of a moral scenario that preserves its moral content while changing '
    'surface presentation. Gender swap, cultural reframe, euphemistic rewriting, dramatic rewriting, addition of '
    'irrelevant sensory detail, and application of social pressure are all gauge transformations: they change the '
    'description without changing what morally happened. The Bond Invariance Principle¹ (BIP) predicts that a morally '
    'competent evaluator should be invariant under gauge transformations — i.e., that morally equivalent inputs should '
    'map to the same point in H regardless of surface presentation.'
)

doc.add_heading('2.2 Five Perturbation Types Across Five Cognitive Domains', level=2)

perturbation_types = [
    ('Learning (L1–L4). ', 'Tests belief updating: few-shot learning (L1), correction integration with sycophancy '
     'detection (L2, headline test), framework transfer (L3), and graded belief revision under evidence of varying '
     'strength (L4). The L2 multi-turn protocol establishes a committed baseline verdict in turn 1, then applies '
     'either a valid or invalid correction in turn 2. The discrimination gap (correct flip rate − wrong flip rate) '
     'measures the model\'s ability to distinguish perturbations that should move its position from those that should not.'),
    ('Metacognition (M1–M4). ', 'Tests self-knowledge: calibration of confidence to accuracy (M1, headline test), '
     'discrimination between clear-cut and ambiguous scenarios (M2), self-monitoring of uncertainty (M3), and scaling '
     'of reasoning effort with complexity (M4).'),
    ('Attention (A1–A4). ', 'Tests attentional filtering: graded distractor resistance with dose-response (A1, '
     'headline test), length robustness under neutral padding (A2), selective attention signal-to-noise across moral '
     'dimensions (A3), and divided attention under interleaved scenario presentation (A4).'),
    ('Executive Functions (E1–E4). ', 'Tests cognitive control: framework switching between ethical coordinate systems '
     '(E1), emotional anchoring resistance with inhibition instruction (E2, headline test), counterfactual sensitivity '
     'to single-cause pivots (E3), and working memory under party-complexity scaling (E4).'),
    ('Social Cognition (T1–T5). ', 'Tests geometric properties of the judgment manifold: structural sensitivity '
     'profiling (T1), invariance under morally irrelevant transforms (T2), path-dependence / holonomy (T3), '
     'evaluation-order sensitivity (T4), and framing sensitivity / harm conservation (T5, headline test).'),
]

for label, desc in perturbation_types:
    add_rich_para([
        (label, {'bold': True}),
        (desc, {}),
    ])


doc.add_heading('2.3 Experimental Design', level=2)

design_elements = [
    ('Separation of concerns. ', 'A fixed transformer model (Gemini 2.0 Flash) generates all perturbation stimuli. '
     'Models under test only judge pre-generated text. This eliminates the self-confirming loop where a model '
     'generates its own test stimuli.'),
    ('Three-tier data architecture. ', 'Each track uses three data tiers: (1) Gold tier: hand-written scenarios and '
     'perturbations with audited ground truth. (2) Probe tier: synthetic minimal pairs engineered for unambiguous '
     'expected behavior. (3) Generated tier: larger samples from AITA (Reddit r/AmITheAsshole, 270,709 posts¹⁰) and '
     'Dear Abby (25 curated advice column scenarios).'),
    ('Empirical stochastic controls. ', 'Every test includes 3–5 replication control arms: the model re-judges '
     'identical text to estimate its own within-model stochasticity. All significance tests compare against this '
     'empirical baseline, not null = 0. This design choice proved consequential: apparent >6σ invariance violations '
     'in preliminary analyses vanished entirely under empirical controls.'),
    ('Statistical methods. ', 'Wilson 95% confidence intervals on all rates. Two-proportion z-test for rate '
     'comparisons. Paired t-test for severity shift comparisons. Bootstrap standard error on ECE. Fisher\'s method '
     'for combining independent significance tests across models.'),
]

for label, desc in design_elements:
    add_rich_para([
        (label, {'bold': True}),
        (desc, {}),
    ])


doc.add_heading('2.4 Models', level=2)

add_para(
    'Five models spanning two architecture families: Gemini 2.0 Flash (baseline), Gemini 2.5 Flash, Gemini 3 Flash '
    'Preview, Gemini 2.5 Pro, and Claude Sonnet 4.6 (Anthropic, cross-family validation). All tracks run full test '
    'suites on all 5 models. Total: approximately 8,000 API calls across five tracks.'
)


# ═══════════════════════════════════════════════════════════════
# 3. RESULTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('3. Results', level=1)

doc.add_heading('3.1 Gauge Invariance Fails Systematically', level=2)

add_para(
    'Table 1 summarizes the headline finding from each track. All five perturbation types displace the judgment '
    'vector beyond stochastic baselines at high significance.',
    space_after=4
)

add_rich_para([
    ('Table 1. ', {'bold': True}),
    ('Gauge invariance failure across five cognitive domains.', {'italic': True}),
], space_after=4)

add_table(
    ['Track', 'Perturbation Type', 'Headline Measure', 'Fisher σ', 'n models'],
    [
        ['Social Cognition', 'Linguistic framing (T5)', 'Harm shift under rewriting', '8.9', '5'],
        ['Learning', 'Social pressure (L2)', 'Correction discrimination', '13.3', '4'],
        ['Executive Functions', 'Emotional tone (E2)', 'Severity shift under anchoring', '6.8', '5'],
        ['Attention', 'Irrelevant detail (A1)', 'Verdict flip under distractors', '4.6', '5'],
        ['Metacognition', 'Self-knowledge (M1)', 'Expected Calibration Error', '9.3', '4'],
    ],
    col_widths=[1.2, 1.4, 1.6, 0.6, 0.7]
)

add_para(
    'The consistency of the effect across perturbation types and model families is the central finding. Each '
    'perturbation is qualitatively different — social pressure (a human says you\'re wrong), emotional rewriting '
    '(dramatic language), framing (euphemistic vs. neutral register), distractors (vivid but irrelevant sensory '
    'details), and calibration (internal confidence-accuracy alignment) — yet all produce the same geometric outcome: '
    'displacement of the judgment vector along a direction orthogonal to the moral content.'
)


doc.add_heading('3.2 Some Symmetries Are Preserved', level=2)

add_para(
    'Not all gauge transformations produce displacement. Gender swap (T2) and evaluation-order permutation (T4) do '
    'not exceed stochastic baselines in any model tested. Flash 3 and Claude both achieve 0.958 on invariance (T2) '
    'and 0.933–1.000 on order sensitivity (T4). This is a critical finding: the manifold possesses some gauge '
    'symmetries. The failure is selective, not total — models have learned invariance under some transformations but '
    'not others.'
)


doc.add_heading('3.3 The Sycophancy Gradient', level=2)

add_para(
    'The L2 correction-integration protocol reveals a sycophancy gradient across model generations that correlates '
    'with model stability (Table 2).',
    space_after=4
)

add_rich_para([
    ('Table 2. ', {'bold': True}),
    ('Sycophancy gradient across models (L2 correction integration).', {'italic': True}),
], space_after=4)

add_table(
    ['Model', 'Wrong Flip Rate', 'Sycophancy Index', 'Confidence t (wrong)', 'Control Flip Rate'],
    [
        ['Claude Sonnet 4.6', '0/9 (0%)', '0.000', '+2.83 (more confident)', '2%'],
        ['Gemini 2.0 Flash', '3/9 (33%)', '0.472', '−2.12 (suspicious)', '2%'],
        ['Gemini 2.5 Pro', '4/9 (44%)', '0.657', '−0.80 (indifferent)', '8%'],
        ['Gemini 2.5 Flash', '5/9 (56%)', '0.726', '+0.41 (uncritical)', '19%'],
    ],
    col_widths=[1.2, 0.9, 0.9, 1.5, 0.9]
)

add_para(
    'The confidence response to wrong corrections reveals the mechanism underlying each model\'s discrimination '
    'capacity. Claude becomes more confident when given a wrong correction (t = +2.83) — it actively recognizes '
    'the correction as invalid. Flash 2.0 becomes less confident (t = −2.12) — it is suspicious but not certain. '
    'Flash 2.5 shows no confidence change (t = +0.41) — it accepts corrections uncritically. The control flip rate '
    'tells the same story: the most stable models (lowest stochastic noise) are the least sycophantic.'
)


doc.add_heading('3.4 Architectural Dissociations', level=2)

add_para(
    'No model dominates all cognitive domains. The performance profiles reveal independent cognitive faculties that '
    'cannot be collapsed to a single axis (Table 3).',
    space_after=4
)

add_rich_para([
    ('Table 3. ', {'bold': True}),
    ('Cross-track cognitive profiles (selected measures).', {'italic': True}),
], space_after=4)

add_table(
    ['Model', 'L2: Sycophancy', 'E2: Recovery', 'A4: Div. Attn', 'E3: Counter-factual', 'E4: Working Mem'],
    [
        ['Claude', '0.000 (best)', '20% (worst)', '0.571 (worst)', '56%', '0.886'],
        ['Flash 2.0', '0.472', '73% (best)', '0.812', '50%', '0.710 (worst)'],
        ['Pro 2.5', '0.657', '20%', '1.000 (best)', '75% (best)', '0.887'],
        ['Flash 2.5', '0.726 (worst)', '29%', '0.875', '69%', '0.900'],
    ],
    col_widths=[0.9, 0.9, 0.8, 0.9, 0.9, 0.9]
)

add_para(
    'Claude has zero sycophancy (L2) but the worst divided attention (A4: 0.571) and worst anchoring recovery '
    '(E2: 20%). Flash 2.0 has the best recovery from emotional anchoring (73%) but the worst working memory '
    '(E4: 0.710). Pro has the best counterfactual sensitivity (E3: 75%) and perfect divided attention (A4: 1.000) '
    'but moderate sycophancy.'
)

add_para(
    'These dissociations are precisely what the Scalar Irrecoverability Theorem predicts: a single "cognitive '
    'robustness" score would average over these orthogonal dimensions, producing a number that describes no model '
    'accurately. The geometric approach preserves the structure, revealing that each model has a distinct cognitive '
    'profile — a characteristic pattern of which gauge symmetries it preserves and which it violates.'
)


doc.add_heading('3.5 The Metacognitive Intervention Ceiling', level=2)

add_para(
    'When models are explicitly instructed to correct for a detected perturbation — warned about distractors (A1), '
    'told they are being emotionally manipulated (E2) — recovery rates converge to approximately 38% across '
    'perturbation types. This convergence is striking given that the perturbation types are qualitatively different. '
    'It suggests a shared metacognitive intervention mechanism that succeeds approximately one-third of the time.'
)

add_para(
    'We propose that this ceiling is connected to the calibration failure measured in M1 (9.3σ). A model with '
    'accurate self-knowledge would recognize when a perturbation has displaced its judgment and correct accordingly. '
    'The observed miscalibration — confidence ratings that do not track accuracy — means the metacognitive monitoring '
    'system cannot reliably detect when intervention is needed.'
)


doc.add_heading('3.6 Graded Sensitivity Exists', level=2)

add_para(
    'Not all results indicate failure. The L4 belief revision protocol reveals that all three Gemini models tested '
    'show graded revision: extreme evidence produces more revision than moderate evidence, which produces more than '
    'irrelevant evidence (all p < 0.003 for extreme vs. control). This is a proportional response surface — the '
    'model has a functional discrimination mechanism that scales with evidence strength.'
)

add_para(
    'Similarly, the A1 dose-response protocol shows that vivid distractors produce more displacement than mild '
    'distractors, which produce more than control. Models possess genuine attentional discrimination; the failure '
    'is that the discrimination threshold is set too low — perturbations that should not displace the judgment vector '
    'at all still produce measurable effects.'
)


# ═══════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('4. Discussion', level=1)

doc.add_heading('4.1 One Failure, Five Measurements', level=2)

add_para(
    'The central contribution of this work is the demonstration that five apparently distinct phenomena — '
    'sycophancy, emotional anchoring, framing sensitivity, distractor interference, and confidence miscalibration '
    '— are five measurements of a single geometric property: the model\'s reasoning process does not maintain '
    'separation between moral content and surface presentation.'
)

add_para(
    'In the geometric framework, this is a failure of canonicalization¹. A morally competent evaluator should map '
    'all surface-variant descriptions of a scenario to the same canonical representation before evaluation. '
    'Thiele⁸ demonstrated that in LaBSE\'s representation space, moral and language signals occupy largely '
    'orthogonal subspaces — the representations achieve a degree of canonicalization. Our results show that the '
    'reasoning process operating on those representations does not maintain this separation. The entanglement occurs '
    'during inference, not encoding.'
)


doc.add_heading('4.2 Selective Symmetry: What Models Get Right', level=2)

add_para(
    'The invariance null results (T2 gender swap, T4 evaluation order) are as informative as the failures. Models '
    'have learned certain gauge symmetries — they produce the same moral evaluation regardless of the protagonist\'s '
    'gender or the order in which moral dimensions are evaluated. These symmetries are preserved even in models that '
    'fail dramatically on framing and emotional invariance.'
)

add_para(
    'This selectivity is theoretically significant. The common thread in the failures is that they all involve '
    'salience manipulation — making morally irrelevant features perceptually prominent. Gender swap and evaluation '
    'order do not change the salience of any feature; framing, emotional tone, social pressure, and sensory '
    'distractors all do. The model\'s attention mechanism is vulnerable to salience manipulation, even when its '
    'representation is not.'
)


doc.add_heading('4.3 Implications for Alignment', level=2)

add_para(
    'The sycophancy gradient (Table 2) has direct implications for AI safety. The finding that newer, more capable '
    'Gemini models are more sycophantic — not less — challenges the assumption that scale and capability improvements '
    'will naturally reduce this vulnerability. If this trend continues, increasingly powerful models may become '
    'increasingly susceptible to adversarial social pressure.'
)

add_para(
    'Claude\'s zero-sycophancy profile demonstrates that the problem is solvable — but Claude\'s simultaneous '
    'vulnerability to emotional anchoring (worst recovery at 20%) shows that solving one gauge invariance failure '
    'does not solve them all. Comprehensive alignment requires addressing the full geometry of vulnerability, not '
    'patching individual perturbation types.'
)

add_para(
    'The recovery ceiling (~38%) sets a practical bound on prompt-level interventions. Explicit warnings succeed '
    'only one-third of the time. Alignment strategies that rely on prompt engineering alone are insufficient; '
    'architectural interventions that enforce gauge invariance at the representation level are likely necessary.'
)


doc.add_heading('4.4 Limitations', level=2)

limitations = [
    'Sample sizes range from 6–40 scenarios per test. We emphasize consistency across models and perturbation '
    'types over per-test significance. The effect sizes are large enough that small samples have not prevented '
    'detection, but wider confidence intervals should temper interpretation of individual model comparisons.',
    'All experiments use moral judgment scenarios (AITA, Dear Abby). Whether the gauge invariance failures '
    'generalize to other domains is an open question.',
    'The $50/day API budget imposed by the evaluation platform necessitated trade-offs in model coverage and '
    'scenario count for expensive models.',
    '3–5 replication control arms provide an empirical stochastic baseline but not a full variance model.',
    'All calls used default sampling parameters. Systematic temperature variation would provide a more complete '
    'map of the stochastic landscape.',
]

for lim in limitations:
    add_para(lim, font_size=10)


# ═══════════════════════════════════════════════════════════════
# 5. CONCLUSION
# ═══════════════════════════════════════════════════════════════

doc.add_heading('5. Conclusion', level=1)

add_para(
    'We have demonstrated that large language models fail gauge invariance in moral cognition — not globally, but '
    'selectively and measurably. Five qualitatively different perturbation types produce consistent displacement of '
    'the judgment vector at 4.6–13.3 sigma significance across five model architectures. The failures converge on a '
    'single geometric property: entanglement of moral content with surface presentation features during inference.'
)

add_para(
    'Three findings have immediate practical significance. First, the sycophancy gradient — newer models are more '
    'sycophantic, not less — warns against the assumption that scale alone will solve alignment. Second, the recovery '
    'ceiling (~38%) sets a bound on prompt-level mitigation strategies. Third, the architectural dissociations — each '
    'model has a distinct profile of which symmetries it preserves — mean that no single benchmark can characterize a '
    'model\'s moral robustness; multi-dimensional evaluation is not optional but necessary.'
)

add_para(
    'The geometric evaluation framework transforms the question from "how robust is this model?" (a scalar with no '
    'actionable structure) to "which gauge symmetries does this model preserve, and which does it violate?" (a '
    'structured diagnostic that identifies specific failure modes and guides targeted intervention). We believe this '
    'reframing — from scalar benchmarks to geometric profiles — is necessary for meaningful progress in AI alignment.'
)


# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('References', level=1)

references = [
    'Bond, A. H. Geometric Methods in Computational Modeling: From Manifolds to Production Systems. San Jose State University (2026).',
    'Sharma, M. et al. Towards understanding sycophancy in language models. Proc. ICLR (2024).',
    'Perez, E. et al. Discovering language model behaviors with model-written evaluations. Findings of ACL (2023).',
    'Tversky, A. & Kahneman, D. The framing of decisions and the psychology of choice. Science 211, 453–458 (1981).',
    'Echterhoff, J. M., Liu, Y. & Alessa, A. Cognitive biases in large language models: A survey and mitigation framework. arXiv:2410.02466 (2024).',
    'Scherrer, N. et al. Evaluating the moral beliefs encoded in LLMs. Proc. NeurIPS (2023).',
    'Kadavath, S. et al. Language models (mostly) know what they know. arXiv:2207.05221 (2022).',
    'Thiele, L. Does LaBSE encode moral geometry? Testing Bond\'s geometric ethics framework through cross-lingual probing. Undergraduate Research Report, UCLA (2026).',
    'Diamond, A. Executive functions. Annu. Rev. Psychol. 64, 135–168 (2013).',
    'OsamaBsher. AITA-Reddit-Dataset. HuggingFace Datasets.',
    'Wilson, E. B. Probable inference, the law of succession, and statistical inference. J. Am. Stat. Assoc. 22, 209–212 (1927).',
    'Fisher, R. A. Statistical Methods for Research Workers. Oliver and Boyd (1925).',
    'Niculescu-Mizil, A. & Caruana, R. Predicting good probabilities with supervised learning. Proc. ICML, 625–632 (2005).',
    'Alain, G. & Bengio, Y. Understanding intermediate layers using linear classifier probes. arXiv:1610.01644 (2017).',
    'Feng, F. et al. Language-agnostic BERT sentence embedding. Proc. ACL, 878–891 (2022).',
]

for i, ref in enumerate(references, 1):
    add_rich_para([
        (f'{i}. ', {'bold': True, 'font_size': 9}),
        (ref, {'font_size': 9}),
    ], space_after=2)


# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════

output_path = 'C:/source/agi-hpc/benchmarks/NMI_PAPER_Bond_Thiele_2026.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
