"""Build NMI paper DOCX from the revised markdown draft.
Uses Garamond for body, Consolas for code/technical terms."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles['Normal']
style.font.name = 'Garamond'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, (size, color, sp_before) in enumerate([
    (18, RGBColor(0x1a, 0x23, 0x32), 24),
    (14, RGBColor(0x2c, 0x3e, 0x50), 18),
    (12, RGBColor(0x34, 0x49, 0x5e), 14),
], start=1):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Garamond'
    h.font.size = Pt(size)
    h.font.bold = True
    h.font.color.rgb = color
    h.paragraph_format.space_before = Pt(sp_before)
    h.paragraph_format.space_after = Pt(8)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Garamond'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1a2332')
        cell.paragraphs[0]._element.get_or_add_pPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Garamond'
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if ri % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'f4f6f8')
                p._element.get_or_add_pPr().append(shading)
    doc.add_paragraph()


def add_rich_text(paragraph, text):
    """Parse simple markdown-like formatting: *italic*, **bold**."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Garamond'
            run.font.size = Pt(11)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Garamond'
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Garamond'
            run.font.size = Pt(11)


# Read the markdown
md = open('C:/source/agi-hpc/benchmarks/NMI_PAPER_DRAFT.md').read()
lines = md.split('\n')

i = 0
in_table = False
table_headers = None
table_rows = []

while i < len(lines):
    line = lines[i]

    # Skip the HR lines
    if line.strip() == '---':
        i += 1
        continue

    # Title
    if line.startswith('# ') and i < 3:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(line[2:].strip())
        run.bold = True
        run.font.name = 'Garamond'
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1a, 0x23, 0x32)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        i += 1
        continue

    # Author line
    if line.startswith('Andrew H. Bond'):
        p = doc.add_paragraph()
        run = p.add_run(line.replace('^1,2^', '¹˒²').replace('^3^', '³').replace('^1^', '¹').replace('^2^', '²'))
        run.font.name = 'Garamond'
        run.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # Affiliation lines
    if line.startswith('^') and ('Department' in line or 'Senior Member' in line):
        p = doc.add_paragraph()
        clean = line.replace('^1^', '¹').replace('^2^', '²').replace('^3^', '³')
        run = p.add_run(clean)
        run.font.name = 'Garamond'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        i += 1
        continue

    # Headings
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue
    if line.startswith('## '):
        title = line[3:].strip()
        level = 2 if title in ['Abstract'] else 1
        doc.add_heading(title, level=level)
        i += 1
        continue

    # Tables
    if '|' in line and line.strip().startswith('|'):
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        # Check if next line is separator
        if i + 1 < len(lines) and '---' in lines[i + 1]:
            # This is a table header
            table_headers = cells
            table_rows = []
            i += 2  # skip header and separator
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                row_cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                table_rows.append(row_cells)
                i += 1
            if table_headers and table_rows:
                add_table(table_headers, table_rows)
            continue
        i += 1
        continue

    # Bold table captions
    if line.startswith('**Table '):
        p = doc.add_paragraph()
        # Extract bold and italic parts
        clean = line.replace('**', '').replace('*', '')
        run = p.add_run(clean)
        run.bold = True
        run.font.name = 'Garamond'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        i += 1
        continue

    # Empty lines
    if not line.strip():
        i += 1
        continue

    # Normal paragraphs
    # Collect continuation lines
    para_text = line
    while i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].startswith('#') and not lines[i + 1].startswith('|') and not lines[i + 1].startswith('**Table') and not lines[i + 1].startswith('---'):
        i += 1
        para_text += ' ' + lines[i]

    # Clean up references
    para_text = para_text.replace('^1^', '¹').replace('^2^', '²').replace('^3^', '³')
    para_text = para_text.replace('^7^', '⁷').replace('^8^', '⁸').replace('^9^', '⁹')
    para_text = para_text.replace('^10^', '¹⁰').replace('^2,3^', '²˒³')
    para_text = para_text.replace('^4^', '⁴').replace('^5^', '⁵').replace('^6^', '⁶')
    para_text = para_text.replace('^5,6^', '⁵˒⁶')

    p = doc.add_paragraph()
    add_rich_text(p, para_text)

    # Abstract gets italic
    if i < 20 and 'Large language models are increasingly' in para_text:
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(10)

    i += 1

# Save
output = 'C:/source/agi-hpc/benchmarks/NMI_PAPER_Bond_Thiele_2026.docx'
doc.save(output)
print(f'Saved revised paper to {output}')
